import librosa
import numpy as np
import matplotlib.pyplot as plt
import os
from docx import Document

def extract_audio_features(audio_file):
    # Load audio file
    y, sr = librosa.load(audio_file, sr=None)
    
    # Extract MFCCs
    mfccs = librosa.feature.mfcc(y=y, sr=sr, n_mfcc=13)
    
    # Extract pitch
    pitch = librosa.yin(y, fmin=librosa.note_to_hz('C1'), fmax=librosa.note_to_hz('C7'))
    
    # Extract intensity
    intensity = np.abs(y)
    
    # Extract spectral centroid
    spectral_centroid = librosa.feature.spectral_centroid(y=y, sr=sr)
    
    # Extract harmonic-to-noise ratio (HNR)
    harmonic, percussive = librosa.effects.hpss(y)
    hnr = librosa.effects.harmonic(y)
    
    return mfccs, pitch, intensity, spectral_centroid, hnr, sr

def plot_features(mfccs, pitch, intensity, spectral_centroid, hnr, sr):
    # Plot MFCCs
    plt.figure(figsize=(10, 6))
    plt.subplot(5, 1, 1)
    librosa.display.specshow(mfccs, sr=sr, x_axis='time')
    plt.colorbar(format='%+2.0f dB')
    plt.title('MFCCs')
    plt.ylabel('MFCC Coefficients')

    # Plot pitch
    plt.subplot(5, 1, 2)
    plt.plot(pitch, label='Pitch')
    plt.title('Pitch')
    plt.ylabel('Frequency (Hz)')

    # Plot intensity
    plt.subplot(5, 1, 3)
    plt.plot(intensity, label='Intensity')
    plt.title('Intensity')
    plt.ylabel('Amplitude')

    # Plot spectral centroid
    plt.subplot(5, 1, 4)
    plt.plot(spectral_centroid.T, label='Spectral Centroid')
    plt.title('Spectral Centroid')
    plt.ylabel('Frequency')

    # Plot HNR
    plt.subplot(5, 1, 5)
    plt.plot(hnr, label='HNR')
    plt.title('Harmonic-to-Noise Ratio (HNR)')
    plt.xlabel('Time')
    plt.ylabel('Amplitude')

    plt.tight_layout()
    plt.show()

def save_features_to_word(mfccs, pitch, intensity, spectral_centroid, hnr, sr, audio_file):
    # Create a new Word document
    doc = Document()
    
    # Add a title
    doc.add_heading(f'Features for {audio_file}', level=1)
    
    # Add a table to store the features
    table = doc.add_table(rows=6, cols=2)
    table.style = 'Table Grid'
    
    # Add feature names and values to the table
    features = [('MFCCs', mfccs.shape), ('Pitch', pitch.shape), ('Intensity', intensity.shape),
                ('Spectral Centroid', spectral_centroid.shape), ('HNR', hnr.shape), ('Sampling Rate', sr)]
    for i, (name, shape) in enumerate(features):
        table.cell(i, 0).text = name
        table.cell(i, 1).text = str(shape)
    
    # Save the Word document
    doc.save(f'features_{os.path.basename(audio_file)}.docx')

# Example usage
audio_file = "output1/cleaned_01_01_01_01_dogs-sitting_fear.wav"  # Replace with the path to your audio file
mfccs, pitch, intensity, spectral_centroid, hnr, sr = extract_audio_features(audio_file)

# Plot the extracted features
plot_features(mfccs, pitch, intensity, spectral_centroid, hnr, sr)

# Save features to Word file
save_features_to_word(mfccs, pitch, intensity, spectral_centroid, hnr, sr, audio_file)
